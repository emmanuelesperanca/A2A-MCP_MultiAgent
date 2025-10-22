"""
API Endpoint para Ingestão de Conhecimento
Baseado em ingest_data/ingest_data.py
"""

from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
from typing import List
import os
import json
import traceback
import tempfile
from dotenv import load_dotenv

# Imports de processamento
import pypdf
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from openai import OpenAI
import psycopg2
import psycopg2.extras

# OCR opcional
try:
    import pytesseract
    import pdf2image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

router = APIRouter()

# =============================================================================
# CONFIGURAÇÕES
# =============================================================================

load_dotenv()  # Carrega variáveis do .env

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
DATABASE_URL = os.getenv("DATABASE_URL")

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
EMBEDDING_MODEL = "text-embedding-3-small"
EMBEDDING_BATCH_SIZE = 100

# Cliente OpenAI
openai_client = OpenAI(api_key=OPENAI_API_KEY)

# =============================================================================
# FUNÇÕES DE EXTRAÇÃO DE TEXTO
# =============================================================================


def extrair_texto_pdf(caminho_arquivo: str) -> str:
    """Extrai texto de PDF"""
    try:
        reader = pypdf.PdfReader(caminho_arquivo)
        texto = "".join(page.extract_text() or "" for page in reader.pages)
        return texto.strip()
    except Exception as e:
        return f"Erro ao ler PDF: {e}"


def extrair_texto_pdf_com_ocr(caminho_arquivo: str) -> str:
    """Extrai texto usando OCR quando PDF é baseado em imagens"""
    import os
    import sys
    from pathlib import Path
    
    if not OCR_AVAILABLE:
        erro_msg = "❌ OCR não disponível - pytesseract ou pdf2image não importado"
        print(erro_msg)
        return f"Erro: {erro_msg}"
    
    try:
        # Debug: verificar configuração do Tesseract
        print(f"\n🔍 === DIAGNÓSTICO OCR ===")
        
        # Detectar se está rodando como executável PyInstaller
        if getattr(sys, 'frozen', False):
            base_path = Path(sys._MEIPASS)
            tesseract_exe = base_path / "tesseract" / "tesseract.exe"
            tessdata_dir = base_path / "tesseract" / "tessdata"
            
            print(f"📁 Base path: {base_path}")
            print(f"📄 Tesseract.exe: {tesseract_exe}")
            print(f"   Existe: {tesseract_exe.exists()}")
            print(f"📁 Tessdata: {tessdata_dir}")
            print(f"   Existe: {tessdata_dir.exists()}")
            
            if tessdata_dir.exists():
                tessdata_files = list(tessdata_dir.glob("*.traineddata"))
                print(f"   Idiomas: {[f.stem for f in tessdata_files]}")
            
            if tesseract_exe.exists():
                pytesseract.pytesseract.tesseract_cmd = str(tesseract_exe)
                # Configurar TESSDATA_PREFIX para apontar PARA A PASTA TESSDATA
                # Tesseract procura os arquivos .traineddata diretamente em $TESSDATA_PREFIX/
                os.environ['TESSDATA_PREFIX'] = str(tessdata_dir)
                print(f"✅ Tesseract configurado!")
                print(f"   TESSDATA_PREFIX: {os.environ.get('TESSDATA_PREFIX')}")
            else:
                erro_msg = f"❌ Tesseract não encontrado em: {tesseract_exe}"
                print(erro_msg)
                return f"Erro: {erro_msg}"
        else:
            print(f"ℹ️ Rodando em modo desenvolvimento")
            print(f"📄 Tesseract cmd: {pytesseract.pytesseract.tesseract_cmd}")
        
        # Converter PDF para imagens (requer Poppler)
        print(f"\n📄 Convertendo PDF para imagens (usando Poppler)...")
        
        # Configurar Poppler se rodando como executável
        poppler_path = None
        if getattr(sys, 'frozen', False):
            base_path = Path(sys._MEIPASS)
            poppler_bin = base_path / "poppler" / "Library" / "bin"
            if poppler_bin.exists():
                poppler_path = str(poppler_bin)
                print(f"✅ Poppler encontrado: {poppler_path}")
            else:
                print(f"⚠️ Poppler não encontrado em: {poppler_bin}")
        
        try:
            images = pdf2image.convert_from_path(
                caminho_arquivo, 
                dpi=300,
                poppler_path=poppler_path
            )
            print(f"✅ {len(images)} páginas convertidas")
        except Exception as poppler_error:
            erro_msg = f"Erro ao converter PDF - Poppler não disponível ou não configurado: {poppler_error}"
            print(f"❌ {erro_msg}")
            print(f"💡 Solução: Baixe o Poppler de https://github.com/oschwartz10612/poppler-windows/releases/")
            print(f"   e coloque em hooks/poppler/ antes de buildar")
            return f"Erro: {erro_msg}"
        
        texto_total = []
        
        for i, image in enumerate(images):
            print(f"   🔍 Processando página {i+1}/{len(images)}...")
            texto_pagina = pytesseract.image_to_string(image, lang='por+eng')
            if texto_pagina.strip():
                texto_total.append(texto_pagina.strip())
                print(f"      ✅ {len(texto_pagina)} caracteres extraídos")
        
        resultado = "\n\n".join(texto_total)
        print(f"✅ OCR concluído: {len(resultado)} caracteres totais")
        return resultado
        
    except Exception as e:
        print(f"❌ ERRO NO OCR: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return f"Erro no OCR: {e}"


def extrair_texto_docx(caminho_arquivo: str) -> str:
    """Extrai texto de DOCX"""
    try:
        document = docx.Document(caminho_arquivo)
        full_text = []
        
        # Parágrafos
        for para in document.paragraphs:
            full_text.append(para.text)
        
        # Tabelas
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        return "\n".join(full_text)
    except Exception as e:
        return f"Erro ao ler DOCX: {e}"


# =============================================================================
# FUNÇÕES DE PROCESSAMENTO
# =============================================================================

def dividir_em_chunks(texto: str) -> List[str]:
    """Divide texto em chunks"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len
    )
    return text_splitter.split_text(texto)


def gerar_embeddings_em_lote(chunks: List[str]) -> List[list]:
    """Gera embeddings para uma lista de chunks"""
    vetores = []
    
    for i in range(0, len(chunks), EMBEDDING_BATCH_SIZE):
        lote = chunks[i:i + EMBEDDING_BATCH_SIZE]
        
        try:
            # Garantir que não há textos vazios
            lote_processado = [text if text.strip() else " " for text in lote]
            
            response = openai_client.embeddings.create(
                input=lote_processado,
                model=EMBEDDING_MODEL
            )
            
            vetores.extend([item.embedding for item in response.data])
        except Exception as e:
            print(f"Erro no lote {i//EMBEDDING_BATCH_SIZE + 1}: {e}")
            # Adicionar None para chunks com erro
            vetores.extend([None] * len(lote))
    
    return vetores


def inserir_chunks_no_db(
    tabela: str,
    chunks: List[str],
    vetores: List[list],
    metadados: dict
) -> int:
    """Insere chunks no banco de dados"""
    conn = psycopg2.connect(DATABASE_URL)
    
    sql = f"""
        INSERT INTO "{tabela}" (
            conteudo_original, fonte_documento, dado_sensivel, apenas_para_si,
            areas_liberadas, nivel_hierarquico_minimo, geografias_liberadas,
            projetos_liberados, idioma, data_validade, responsavel, aprovador, vetor
        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s);
    """
    
    dados_para_inserir = []
    
    for chunk, vetor in zip(chunks, vetores):
        if vetor is not None:
            dados_para_inserir.append((
                chunk,
                metadados['fonte_documento'],
                metadados['dado_sensivel'],
                metadados['apenas_para_si'],
                metadados['areas_liberadas'],
                metadados['nivel_hierarquico_minimo'],
                metadados['geografias_liberadas'],
                metadados['projetos_liberados'],
                metadados['idioma'],
                metadados['data_validade'],
                metadados['responsavel'],
                metadados['aprovador'],
                vetor
            ))
    
    if not dados_para_inserir:
        conn.close()
        return 0
    
    try:
        with conn.cursor() as cur:
            psycopg2.extras.execute_batch(cur, sql, dados_para_inserir)
        conn.commit()
        count = len(dados_para_inserir)
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        conn.close()
    
    return count


# =============================================================================
# ENDPOINT
# =============================================================================

@router.post("/api/knowledge/ingest")
async def ingest_knowledge(
    file: UploadFile = File(...),
    target_table: str = Form(...),
    metadata: str = Form(...)
):
    """
    Endpoint para ingestão de conhecimento
    
    Args:
        file: Arquivo PDF ou DOCX
        target_table: Tabela de destino no banco
        metadata: JSON com metadados de governança
    """
    
    try:
        # Parse metadata
        meta = json.loads(metadata)
        
        # Salvar arquivo temporariamente
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            # Extrair texto
            print(f"\n📄 Processando arquivo: {file.filename}")
            
            if file.filename.lower().endswith('.pdf'):
                print("   → Tentando extração normal de PDF...")
                texto = extrair_texto_pdf(tmp_path)
                
                # Tentar OCR se necessário
                if not texto or len(texto.strip()) < 50:
                    print("   ⚠️ Pouco texto extraído, tentando OCR...")
                    texto_ocr = extrair_texto_pdf_com_ocr(tmp_path)
                    
                    if texto_ocr and not texto_ocr.startswith("Erro"):
                        texto = texto_ocr
                        print("   ✅ OCR bem-sucedido!")
                    else:
                        print(f"   ❌ OCR falhou: {texto_ocr}")
                    
            elif file.filename.lower().endswith('.docx'):
                print("   → Extraindo texto de DOCX...")
                texto = extrair_texto_docx(tmp_path)
            else:
                raise HTTPException(status_code=400, detail="Formato não suportado")
            
            # Verificar se há erros explícitos
            if texto.startswith("Erro"):
                print(f"   ❌ Erro na extração: {texto}")
                raise HTTPException(status_code=400, detail=f"Erro ao processar arquivo: {texto}")
            
            if not texto.strip():
                print("   ❌ Nenhum texto extraído")
                raise HTTPException(status_code=400, detail="Não foi possível extrair texto do arquivo")
            
            print(f"   ✅ Texto extraído: {len(texto)} caracteres")
            
            # Limpar texto
            texto_limpo = "\n\n".join([
                linha.strip()
                for linha in texto.replace('\xa0', ' ').splitlines()
                if linha.strip()
            ])
            
            # Dividir em chunks
            chunks = dividir_em_chunks(texto_limpo)
            
            if not chunks:
                raise HTTPException(status_code=400, detail="Nenhum chunk gerado")
            
            # Gerar embeddings
            vetores = gerar_embeddings_em_lote(chunks)
            
            # Filtrar chunks válidos
            chunks_validos = [chunk for chunk, vetor in zip(chunks, vetores) if vetor is not None]
            vetores_validos = [vetor for vetor in vetores if vetor is not None]
            
            if not vetores_validos:
                raise HTTPException(status_code=500, detail="Nenhum embedding gerado")
            
            # Inserir no banco
            inserted = inserir_chunks_no_db(
                target_table,
                chunks_validos,
                vetores_validos,
                meta
            )
            
            return JSONResponse({
                "success": True,
                "filename": file.filename,
                "text_length": len(texto_limpo),
                "chunks_count": len(chunks),
                "embeddings_count": len(vetores_validos),
                "inserted_count": inserted
            })
            
        finally:
            # Limpar arquivo temporário
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"Erro na ingestão: {e}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

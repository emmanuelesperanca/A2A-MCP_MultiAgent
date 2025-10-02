import os
import sys
import re
import psycopg2
import psycopg2.extras
import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox, Toplevel
import threading
import queue
import traceback
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
import pypdf
import docx

# OCR Dependencies
try:
    import pytesseract
    import pdf2image
    from PIL import Image

    tesseract_paths = [
        r"C:\Users\u137147\AppData\Local\Programs\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
    ]

    for path in tesseract_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            break

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# =============================================================================
# SEÇÃO DE CONFIGURAÇÃO DE PADRÕES
# =============================================================================
AREAS_PADRAO = ["RH", "TI", "Financeiro", "Juridico", "Vendas", "Marketing", "Operações", "ALL"]
GEOGRAFIAS_PADRAO = ["BR", "US", "AR", "CL", "CO", "MX", "ALL_LATAM", "ALL_EUROPE", "ALL"]
PROJETOS_PADRAO = ["N/A"]
IDIOMAS_PADRAO = ["pt-br", "en-us", "es-es"]

NIVEIS_HIERARQUICOS = {
    "1: Estagiário/Aprendiz": 1, "2: Analista Júnior": 2, "3: Analista Pleno/Sênior": 3,
    "4: Coordenador/Gerente": 4, "5: Diretor/Executivo": 5
}

TABELAS_AGENTE = {
    "Agente RH": "knowledge_rh",
    "Agente Tech": "knowledge_tech"
}

# =============================================================================
# CONFIGURAÇÃO E LÓGICA DE NEGÓCIO
# =============================================================================
load_dotenv()

try:
    if not os.getenv("OPENAI_API_KEY"):
        raise ValueError("A chave OPENAI_API_KEY não foi encontrada ou está vazia no arquivo .env.")
    openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
except Exception as e:
    root = tk.Tk(); root.withdraw()
    messagebox.showerror("Erro Crítico de Configuração", f"Não foi possível inicializar a aplicação.\n\nDetalhes: {e}")
    sys.exit(1)

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
EMBEDDING_MODEL = "text-embedding-3-small"
EMBEDDING_BATCH_SIZE = 100

def extrair_texto_pdf(caminho_arquivo):
    """Extrai texto de PDF, com fallback para OCR se necessário."""
    try:
        reader = pypdf.PdfReader(caminho_arquivo)
        texto = "".join(page.extract_text() for page in reader.pages if page.extract_text())
        return texto
    except Exception as e:
        return f"Erro ao ler o PDF: {e}"


def extrair_texto_pdf_com_ocr(caminho_arquivo, log_queue=None):
    """Extrai texto usando OCR quando PDF é baseado em imagens."""
    if not OCR_AVAILABLE:
        erro = "OCR não disponível. Instale: pip install pytesseract pdf2image pillow"
        if log_queue:
            log_queue.put(f"ERRO OCR: {erro}")
        return erro

    try:
        if log_queue:
            log_queue.put("  -> Convertendo PDF para imagens para OCR...")

        # Converter PDF para imagens
        images = pdf2image.convert_from_path(caminho_arquivo, dpi=300)

        if log_queue:
            log_queue.put(f"  -> {len(images)} páginas convertidas, executando OCR...")

        texto_total = []
        for i, image in enumerate(images):
            if log_queue:
                log_queue.put(f"  -> Processando página {i+1}/{len(images)} com OCR...")

            # Aplicar OCR na imagem
            texto_pagina = pytesseract.image_to_string(image, lang='por+eng')
            if texto_pagina.strip():
                texto_total.append(texto_pagina.strip())

        resultado = "\n\n".join(texto_total)

        if log_queue:
            log_queue.put(f"  -> OCR concluído: {len(resultado)} caracteres extraídos")

        return resultado if resultado.strip() else "Nenhum texto extraído via OCR"

    except Exception as e:
        erro = f"Erro no processamento OCR: {e}"
        if log_queue:
            log_queue.put(f"ERRO OCR: {erro}")
        return erro


def extrair_texto_docx(caminho_arquivo):
    try:
        document = docx.Document(caminho_arquivo)
        full_text = [para.text for para in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)
    except Exception as e:
        return f"Erro ao ler o DOCX: {e}"


def dividir_em_chunks(texto):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP, length_function=len)
    return text_splitter.split_text(texto)


def gerar_embeddings_em_lote(chunks, log_queue):
    vetores = []
    for i in range(0, len(chunks), EMBEDDING_BATCH_SIZE):
        lote = chunks[i:i + EMBEDDING_BATCH_SIZE]
        log_queue.put(f"  -> Gerando embeddings para o lote {i//EMBEDDING_BATCH_SIZE + 1}...")
        try:
            lote_processado = [text if text.strip() else " " for text in lote]
            response = openai_client.embeddings.create(input=lote_processado, model=EMBEDDING_MODEL)
            vetores.extend([item.embedding for item in response.data])
        except Exception as e:
            log_queue.put(f"  -> ERRO no lote {i//EMBEDDING_BATCH_SIZE + 1}: {e}")
            vetores.extend([None] * len(lote))
    return vetores


def inserir_chunks_em_lote_no_db(conn, nome_tabela, chunks, metadados, vetores):
    sql = f"""
        INSERT INTO {nome_tabela} (
            conteudo_original, fonte_documento, dado_sensivel, apenas_para_si,
            areas_liberadas, nivel_hierarquico_minimo, geografias_liberadas,
            projetos_liberados, idioma, data_validade, responsavel, aprovador, vetor
        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s);
    """
    dados_para_inserir = []
    for chunk, vetor in zip(chunks, vetores):
        if vetor is not None:
            dados_para_inserir.append((
                chunk, metadados['fonte_documento'], metadados['dado_sensivel'], metadados['apenas_para_si'],
                metadados['areas_liberadas'], metadados['nivel_hierarquico_minimo'], metadados['geografias_liberadas'],
                metadados['projetos_liberados'], metadados['idioma'], metadados['data_validade'] or None,
                metadados['responsavel'], metadados['aprovador'], vetor
            ))
    if not dados_para_inserir:
        return "Nenhum dado válido para inserir."
    try:
        with conn.cursor() as cur:
            psycopg2.extras.execute_batch(cur, sql, dados_para_inserir)
        conn.commit()
        return len(dados_para_inserir)
    except Exception as e:
        conn.rollback()
        return f"Erro ao inserir em lote no banco de dados: {e}"


def conectar_db():
    try:
        return psycopg2.connect(os.getenv("DATABASE_URL"))
    except Exception as e:
        return f"Erro ao conectar ao banco de dados: {e}"


class IngestionApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Ferramenta de Ingestão de Dados - Agente IA v2.13 (Tabela Padronizada)")
        self.root.geometry("850x950")
        self.log_queue = queue.Queue()
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        self.create_widgets(main_frame)
        self.process_log_queue()

    def create_widgets(self, parent):
        file_frame = ttk.LabelFrame(parent, text="1. Seleção de Documentos", padding="10")
        file_frame.pack(fill=tk.X, expand=True, pady=5)

        ttk.Label(file_frame, text="Arquivos:").grid(row=0, column=0, sticky="nw", padx=5, pady=5)
        self.filepath_text = scrolledtext.ScrolledText(file_frame, wrap=tk.NONE, height=5, width=70)
        self.filepath_text.grid(row=0, column=1, sticky="ew", padx=5, pady=5)
        self.filepath_text.configure(state='disabled')
        ttk.Button(file_frame, text="Procurar...", command=self.browse_files).grid(row=0, column=2, sticky="ne", padx=5, pady=5)

        # Alterado para dropdown
        self.tablename_var = tk.StringVar()
        ttk.Label(file_frame, text="Tabela de Destino:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        tablename_combo = ttk.Combobox(file_frame, textvariable=self.tablename_var, values=list(TABELAS_AGENTE.keys()), state='readonly', width=68)
        tablename_combo.grid(row=1, column=1, columnspan=2, sticky="ew", padx=5, pady=5)
        file_frame.columnconfigure(1, weight=1)

        metadata_frame = ttk.LabelFrame(parent, text="2. Metadados de Governança", padding="10")
        metadata_frame.pack(fill=tk.X, expand=True, pady=5)
        self.entries = {}

        ttk.Label(metadata_frame, text="Responsável:").grid(row=1, column=0, sticky="w", padx=5, pady=3)
        responsavel_var = tk.StringVar()
        ttk.Entry(metadata_frame, textvariable=responsavel_var, width=50).grid(row=1, column=1, sticky="ew", padx=5, pady=3)
        self.entries['responsavel'] = responsavel_var
        try:
            self.entries['responsavel'].set(os.getlogin())
        except Exception:
            pass
        ttk.Label(metadata_frame, text="Aprovador:").grid(row=2, column=0, sticky="w", padx=5, pady=3)
        aprovador_var = tk.StringVar()
        ttk.Entry(metadata_frame, textvariable=aprovador_var, width=50).grid(row=2, column=1, sticky="ew", padx=5, pady=3)
        self.entries['aprovador'] = aprovador_var

        self.areas_var = tk.StringVar()
        self.create_multi_select(metadata_frame, "Áreas Liberadas:", self.areas_var, AREAS_PADRAO, 3)
        self.entries['areas_liberadas'] = self.areas_var

        self.geografias_var = tk.StringVar()
        self.create_multi_select(metadata_frame, "Geografias Liberadas:", self.geografias_var, GEOGRAFIAS_PADRAO, 4)
        self.entries['geografias_liberadas'] = self.geografias_var

        self.projetos_var = tk.StringVar()
        self.create_multi_select(metadata_frame, "Projetos Liberados:", self.projetos_var, PROJETOS_PADRAO, 5)
        self.entries['projetos_liberados'] = self.projetos_var

        ttk.Label(metadata_frame, text="Nível Hierárquico Mínimo:").grid(row=6, column=0, sticky="w", padx=5, pady=3)
        self.nivel_var = tk.StringVar()
        nivel_combo = ttk.Combobox(metadata_frame, textvariable=self.nivel_var, values=list(NIVEIS_HIERARQUICOS.keys()), state='readonly')
        nivel_combo.grid(row=6, column=1, columnspan=2, sticky="ew", padx=5, pady=3)
        self.entries['nivel_hierarquico_minimo'] = self.nivel_var

        ttk.Label(metadata_frame, text="Idioma:").grid(row=7, column=0, sticky="w", padx=5, pady=3)
        self.idioma_var = tk.StringVar()
        idioma_combo = ttk.Combobox(metadata_frame, textvariable=self.idioma_var, values=IDIOMAS_PADRAO, state='readonly')
        idioma_combo.grid(row=7, column=1, columnspan=2, sticky="ew", padx=5, pady=3)
        self.entries['idioma'] = self.idioma_var

        ttk.Label(metadata_frame, text="Data de Validade:").grid(row=8, column=0, sticky="w", padx=5, pady=3)
        self.data_var = tk.StringVar()
        self.data_entry = ttk.Entry(metadata_frame, textvariable=self.data_var, width=50)
        self.data_entry.grid(row=8, column=1, columnspan=2, sticky="ew", padx=5, pady=3)
        self.entries['data_validade'] = self.data_var
        self.add_placeholder(self.data_entry, "AAAA-MM-DD")

        self.dado_sensivel_var = tk.BooleanVar()
        self.apenas_para_si_var = tk.BooleanVar()
        ttk.Checkbutton(metadata_frame, text="Dado Sensível", variable=self.dado_sensivel_var).grid(row=9, column=0, sticky="w", padx=5, pady=5)
        ttk.Checkbutton(metadata_frame, text="Apenas para Si", variable=self.apenas_para_si_var).grid(row=9, column=1, sticky="w", padx=5, pady=5)
        metadata_frame.columnconfigure(1, weight=1)

        action_frame = ttk.LabelFrame(parent, text="3. Execução e Log", padding="10")
        action_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        self.start_button = ttk.Button(action_frame, text="Iniciar Ingestão", command=self.start_ingestion_thread)
        self.start_button.pack(pady=10)
        self.log_area = scrolledtext.ScrolledText(action_frame, wrap=tk.WORD, height=15)
        self.log_area.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.log_area.configure(state='disabled')

    def add_placeholder(self, entry, placeholder):
        entry.insert(0, placeholder)
        entry.config(foreground="grey")
        entry.bind("<FocusIn>", lambda event, e=entry, p=placeholder: self.remove_placeholder(e, p))
        entry.bind("<FocusOut>", lambda event, e=entry, p=placeholder: self.set_placeholder(e, p))

    def remove_placeholder(self, entry, placeholder):
        if entry.get() == placeholder:
            entry.delete(0, tk.END)
            entry.config(foreground="black")

    def set_placeholder(self, entry, placeholder):
        if not entry.get():
            entry.insert(0, placeholder)
            entry.config(foreground="grey")

    def create_multi_select(self, parent, label_text, string_var, options, row_num):
        ttk.Label(parent, text=label_text).grid(row=row_num, column=0, sticky="w", padx=5, pady=3)
        entry = ttk.Entry(parent, textvariable=string_var, state='readonly')
        entry.grid(row=row_num, column=1, sticky="ew", padx=5, pady=3)
        button = ttk.Button(parent, text="Selecionar...", command=lambda: self.open_multi_select_dialog(label_text, options, string_var))
        button.grid(row=row_num, column=2, sticky="ew", padx=5, pady=3)

    def open_multi_select_dialog(self, title, options, string_var):
        dialog = Toplevel(self.root)
        dialog.title(f"Selecionar {title.replace(':', '')}")
        dialog.geometry("350x450")
        dialog.transient(self.root); dialog.grab_set()
        listbox_frame = ttk.Frame(dialog, padding="10")
        listbox_frame.pack(fill=tk.BOTH, expand=True)
        listbox = tk.Listbox(listbox_frame, selectmode=tk.MULTIPLE, exportselection=False)
        listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar = ttk.Scrollbar(listbox_frame, orient=tk.VERTICAL, command=listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        listbox.config(yscrollcommand=scrollbar.set)
        for option in options:
            listbox.insert(tk.END, option)
        current_selections = [item.strip() for item in string_var.get().split(',') if item.strip()]
        for i, option in enumerate(options):
            if option in current_selections:
                listbox.selection_set(i)

        def on_ok():
            selected_indices = listbox.curselection()
            selected_values = [listbox.get(i) for i in selected_indices]
            string_var.set(",".join(selected_values))
            dialog.destroy()
        ok_button = ttk.Button(dialog, text="OK", command=on_ok)
        ok_button.pack(pady=10)

    def browse_files(self):
        filepaths = filedialog.askopenfilenames(filetypes=(("Documentos Suportados", "*.docx *.pdf"), ("Todos", "*.*")))
        if filepaths:
            self.filepath_text.configure(state='normal')
            self.filepath_text.delete("1.0", tk.END)
            self.filepath_text.insert("1.0", "\n".join(filepaths))
            self.filepath_text.configure(state='disabled')

    def start_ingestion_thread(self):
        filepaths_str = self.filepath_text.get("1.0", tk.END).strip()
        selected_agent_display_name = self.tablename_var.get()

        if not filepaths_str or not selected_agent_display_name:
            self.log("ERRO: Selecione um ou mais arquivos e especifique uma Tabela de Destino.")
            return

        # Validação da Data
        date_str = self.data_var.get()
        if date_str and date_str != "AAAA-MM-DD":
            try:
                datetime.strptime(date_str, "%Y-%m-%d")
            except ValueError:
                messagebox.showwarning("Formato Inválido", "O formato da Data de Validade está incorreto. Por favor, use AAAA-MM-DD ou deixe o campo em branco.")
                return

        actual_tablename = TABELAS_AGENTE.get(selected_agent_display_name)

        filepaths = filepaths_str.split('\n')
        metadata_template = {key: var.get() for key, var in self.entries.items()}
        nivel_display = metadata_template.pop('nivel_hierarquico_minimo')
        metadata_template['nivel_hierarquico_minimo'] = NIVEIS_HIERARQUICOS.get(nivel_display, 1)
        metadata_template["dado_sensivel"] = self.dado_sensivel_var.get()
        metadata_template["apenas_para_si"] = self.apenas_para_si_var.get()
        for key in ["areas_liberadas", "geografias_liberadas", "projetos_liberados"]:
            metadata_template[key] = [item.strip() for item in metadata_template[key].split(',') if item.strip()] if metadata_template[key] else None

        if metadata_template['data_validade'] == "AAAA-MM-DD":
            metadata_template['data_validade'] = ''

        self.start_button.config(state="disabled"); self.log("Iniciando ingestão em lote...")
        threading.Thread(target=self.run_ingestion, args=(filepaths, actual_tablename, metadata_template), daemon=True).start()

    def log(self, message):
        self.log_area.configure(state='normal'); self.log_area.insert(tk.END, message + "\n")
        self.log_area.configure(state='disabled'); self.log_area.see(tk.END)

    def process_log_queue(self):
        try:
            message = self.log_queue.get_nowait()
            self.log(message)
        except queue.Empty:
            pass
        self.root.after(100, self.process_log_queue)

    def run_ingestion(self, filepaths, tablename, metadata_template):
        total_sucessos = 0
        total_chunks_gerados = 0
        conn = None
        try:
            conn = conectar_db()
            if isinstance(conn, str): self.log_queue.put(conn); return
            for i, filepath in enumerate(filepaths):
                try:
                    self.log_queue.put(f"\n--- Processando Arquivo {i+1}/{len(filepaths)}: {os.path.basename(filepath)} ---")
                    if filepath.lower().endswith('.pdf'):
                        texto_bruto = extrair_texto_pdf(filepath)
                        if (not texto_bruto or
                            not texto_bruto.strip() or
                            "Erro" in texto_bruto or
                            len(texto_bruto.strip()) < 50):
                            self.log_queue.put("  -> PDF com pouco/nenhum texto extraível. Tentando OCR...")
                            texto_bruto = extrair_texto_pdf_com_ocr(filepath, self.log_queue)
                    elif filepath.lower().endswith('.docx'):
                        texto_bruto = extrair_texto_docx(filepath)
                    else:
                        self.log_queue.put("AVISO: Formato não suportado. Pulando.")
                        continue

                    if "Erro" in texto_bruto:
                        self.log_queue.put(texto_bruto)
                        continue

                    if not texto_bruto or not texto_bruto.strip():
                        self.log_queue.put("AVISO: Nenhum texto extraído mesmo com OCR. Pulando.")
                        continue
                    texto_limpo = "\n\n".join([linha.strip() for linha in texto_bruto.replace('\xa0', ' ').splitlines() if linha.strip()])
                    self.log_queue.put(f"Texto extraído e limpo ({len(texto_limpo)} caracteres).")
                    chunks = dividir_em_chunks(texto_limpo)
                    if not chunks:
                        self.log_queue.put("AVISO: Nenhum chunk gerado. Pulando.")
                        continue
                    self.log_queue.put(f"Dividido em {len(chunks)} chunks.")
                    total_chunks_gerados += len(chunks)
                    vetores = gerar_embeddings_em_lote(chunks, self.log_queue)
                    chunks_validos = [chunk for chunk, vetor in zip(chunks, vetores) if vetor is not None]
                    vetores_validos = [vetor for vetor in vetores if vetor is not None]
                    self.log_queue.put(f"  -> {len(vetores_validos)} embeddings gerados com sucesso.")
                    if not vetores_validos:
                        self.log_queue.put("AVISO: Nenhum embedding gerado. Pulando.")
                        continue
                    metadata = metadata_template.copy()
                    metadata['fonte_documento'] = os.path.basename(filepath)
                    self.log_queue.put(f"Inserindo {len(chunks_validos)} chunks no banco de dados...")
                    resultado = inserir_chunks_em_lote_no_db(conn, tablename, chunks_validos, metadata, vetores_validos)
                    if isinstance(resultado, int):
                        self.log_queue.put(f"  -> {resultado} chunks inseridos com sucesso.")
                        total_sucessos += resultado
                    else:
                        self.log_queue.put(f"  -> ERRO na inserção: {resultado}")
                except Exception as e_file:
                    self.log_queue.put(f"!! ERRO CRÍTICO ao processar {os.path.basename(filepath)}: {e_file}")
                    self.log_queue.put(traceback.format_exc())
        except Exception as e_main:
            self.log_queue.put(f"!! ERRO CRÍTICO NO PROCESSO DE INGESTÃO: {e_main}")
            self.log_queue.put(traceback.format_exc())
        finally:
            if conn:
                conn.close()
            self.log_queue.put("\n==============================================")
            self.log_queue.put("Processo de ingestão em lote finalizado.")
            self.log_queue.put(f"Total de Chunks Inseridos: {total_sucessos} de {total_chunks_gerados} gerados.")
            self.start_button.config(state="normal")


if __name__ == "__main__":
    root = tk.Tk()
    style = ttk.Style(root)
    style.theme_use("clam")
    app = IngestionApp(root)
    root.mainloop()
